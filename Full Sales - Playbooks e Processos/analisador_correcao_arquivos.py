#!/usr/bin/env python3
"""
Script para analisar arquivos corrompidos e converter XLSX para CSV
Autor: Assistente AI
Data: 2025-01-02
"""

import os
import sys
import logging
import shutil
import zipfile
from pathlib import Path
from datetime import datetime
import json

# Importações para diferentes tipos de arquivo
try:
    from docx import Document
    import docx
except ImportError:
    print("Instalando python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    import docx

try:
    import pandas as pd
except ImportError:
    print("Instalando pandas...")
    os.system("pip install pandas")
    import pandas as pd

try:
    from openpyxl import load_workbook
except ImportError:
    print("Instalando openpyxl...")
    os.system("pip install openpyxl")
    from openpyxl import load_workbook

try:
    import PyPDF2
except ImportError:
    print("Instalando PyPDF2...")
    os.system("pip install PyPDF2")
    import PyPDF2

try:
    from pptx import Presentation
except ImportError:
    print("Instalando python-pptx...")
    os.system("pip install python-pptx")
    from pptx import Presentation

class AnalisadorCorrecaoArquivos:
    def __init__(self, diretorio_consolidado):
        self.diretorio_consolidado = Path(diretorio_consolidado)
        self.diretorio_corrigidos = self.diretorio_consolidado / "arquivos_corrigidos"
        self.diretorio_csv = self.diretorio_consolidado / "arquivos_csv"
        self.diretorio_logs = self.diretorio_consolidado / "logs"
        
        # Criar diretórios se não existirem
        self.diretorio_corrigidos.mkdir(exist_ok=True)
        self.diretorio_csv.mkdir(exist_ok=True)
        self.diretorio_logs.mkdir(exist_ok=True)
        
        # Configurar logging
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        log_file = self.diretorio_logs / f"analise_correcao_{timestamp}.log"
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file, encoding='utf-8'),
                logging.StreamHandler(sys.stdout)
            ]
        )
        self.logger = logging.getLogger(__name__)
        
        # Estatísticas
        self.stats = {
            'total_arquivos': 0,
            'arquivos_corrompidos': 0,
            'arquivos_corrigidos': 0,
            'xlsx_convertidos': 0,
            'erros': []
        }
        
        self.arquivos_corrompidos = []
        self.arquivos_xlsx = []

    def verificar_arquivo_docx(self, caminho_arquivo):
        """Verifica se um arquivo DOCX está corrompido"""
        try:
            doc = Document(caminho_arquivo)
            # Tenta acessar propriedades básicas
            paragrafos = len(doc.paragraphs)
            self.logger.info(f"DOCX OK: {caminho_arquivo.name} ({paragrafos} parágrafos)")
            return True, None
        except Exception as e:
            self.logger.error(f"DOCX CORROMPIDO: {caminho_arquivo.name} - {str(e)}")
            return False, str(e)

    def verificar_arquivo_xlsx(self, caminho_arquivo):
        """Verifica se um arquivo XLSX está corrompido"""
        try:
            workbook = load_workbook(caminho_arquivo)
            sheets = len(workbook.sheetnames)
            self.logger.info(f"XLSX OK: {caminho_arquivo.name} ({sheets} planilhas)")
            return True, None
        except Exception as e:
            self.logger.error(f"XLSX CORROMPIDO: {caminho_arquivo.name} - {str(e)}")
            return False, str(e)

    def verificar_arquivo_pdf(self, caminho_arquivo):
        """Verifica se um arquivo PDF está corrompido"""
        try:
            with open(caminho_arquivo, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                pages = len(reader.pages)
                self.logger.info(f"PDF OK: {caminho_arquivo.name} ({pages} páginas)")
                return True, None
        except Exception as e:
            self.logger.error(f"PDF CORROMPIDO: {caminho_arquivo.name} - {str(e)}")
            return False, str(e)

    def verificar_arquivo_pptx(self, caminho_arquivo):
        """Verifica se um arquivo PPTX está corrompido"""
        try:
            prs = Presentation(caminho_arquivo)
            slides = len(prs.slides)
            self.logger.info(f"PPTX OK: {caminho_arquivo.name} ({slides} slides)")
            return True, None
        except Exception as e:
            self.logger.error(f"PPTX CORROMPIDO: {caminho_arquivo.name} - {str(e)}")
            return False, str(e)

    def verificar_arquivo_zip(self, caminho_arquivo):
        """Verifica se um arquivo ZIP (base dos Office) está corrompido"""
        try:
            with zipfile.ZipFile(caminho_arquivo, 'r') as zip_file:
                # Testa a integridade do ZIP
                bad_file = zip_file.testzip()
                if bad_file:
                    return False, f"Arquivo corrompido no ZIP: {bad_file}"
                return True, None
        except Exception as e:
            return False, str(e)

    def tentar_reparar_docx(self, caminho_arquivo):
        """Tenta reparar um arquivo DOCX corrompido"""
        try:
            # Método 1: Tentar extrair e recriar
            nome_arquivo = caminho_arquivo.stem
            arquivo_reparado = self.diretorio_corrigidos / f"{nome_arquivo}_reparado.docx"
            
            # Verifica se é um ZIP válido primeiro
            is_zip_ok, zip_error = self.verificar_arquivo_zip(caminho_arquivo)
            
            if is_zip_ok:
                # Tenta abrir como documento e recriar
                try:
                    doc = Document(caminho_arquivo)
                    novo_doc = Document()
                    
                    # Copia parágrafos
                    for para in doc.paragraphs:
                        novo_doc.add_paragraph(para.text)
                    
                    novo_doc.save(arquivo_reparado)
                    self.logger.info(f"DOCX reparado: {arquivo_reparado}")
                    return True, arquivo_reparado
                except:
                    pass
            
            # Método 2: Tentar extrair texto usando python-docx2txt
            try:
                import docx2txt
                texto = docx2txt.process(caminho_arquivo)
                if texto:
                    novo_doc = Document()
                    novo_doc.add_paragraph(texto)
                    novo_doc.save(arquivo_reparado)
                    self.logger.info(f"DOCX reparado (texto extraído): {arquivo_reparado}")
                    return True, arquivo_reparado
            except ImportError:
                os.system("pip install docx2txt")
                import docx2txt
                texto = docx2txt.process(caminho_arquivo)
                if texto:
                    novo_doc = Document()
                    novo_doc.add_paragraph(texto)
                    novo_doc.save(arquivo_reparado)
                    self.logger.info(f"DOCX reparado (texto extraído): {arquivo_reparado}")
                    return True, arquivo_reparado
            except:
                pass
                
            return False, "Não foi possível reparar o arquivo DOCX"
            
        except Exception as e:
            return False, str(e)

    def converter_xlsx_para_csv(self, caminho_arquivo):
        """Converte arquivo XLSX para CSV"""
        try:
            nome_arquivo = caminho_arquivo.stem
            
            # Lê o arquivo Excel
            excel_file = pd.ExcelFile(caminho_arquivo)
            
            # Se tem múltiplas planilhas, cria um CSV para cada
            if len(excel_file.sheet_names) > 1:
                for i, sheet_name in enumerate(excel_file.sheet_names):
                    df = pd.read_excel(caminho_arquivo, sheet_name=sheet_name)
                    csv_file = self.diretorio_csv / f"{nome_arquivo}_planilha_{i+1}_{sheet_name}.csv"
                    df.to_csv(csv_file, index=False, encoding='utf-8-sig')
                    self.logger.info(f"CSV criado: {csv_file}")
            else:
                # Uma única planilha
                df = pd.read_excel(caminho_arquivo)
                csv_file = self.diretorio_csv / f"{nome_arquivo}.csv"
                df.to_csv(csv_file, index=False, encoding='utf-8-sig')
                self.logger.info(f"CSV criado: {csv_file}")
            
            return True, "Conversão realizada com sucesso"
            
        except Exception as e:
            self.logger.error(f"Erro ao converter XLSX para CSV: {caminho_arquivo.name} - {str(e)}")
            return False, str(e)

    def analisar_arquivo(self, caminho_arquivo):
        """Analisa um arquivo específico"""
        extensao = caminho_arquivo.suffix.lower()
        
        if extensao == '.docx':
            is_ok, error = self.verificar_arquivo_docx(caminho_arquivo)
            if not is_ok:
                self.arquivos_corrompidos.append({
                    'arquivo': str(caminho_arquivo),
                    'tipo': 'DOCX',
                    'erro': error
                })
                self.stats['arquivos_corrompidos'] += 1
        
        elif extensao == '.xlsx':
            is_ok, error = self.verificar_arquivo_xlsx(caminho_arquivo)
            if is_ok:
                self.arquivos_xlsx.append(caminho_arquivo)
            else:
                self.arquivos_corrompidos.append({
                    'arquivo': str(caminho_arquivo),
                    'tipo': 'XLSX',
                    'erro': error
                })
                self.stats['arquivos_corrompidos'] += 1
        
        elif extensao == '.pdf':
            is_ok, error = self.verificar_arquivo_pdf(caminho_arquivo)
            if not is_ok:
                self.arquivos_corrompidos.append({
                    'arquivo': str(caminho_arquivo),
                    'tipo': 'PDF',
                    'erro': error
                })
                self.stats['arquivos_corrompidos'] += 1
        
        elif extensao == '.pptx':
            is_ok, error = self.verificar_arquivo_pptx(caminho_arquivo)
            if not is_ok:
                self.arquivos_corrompidos.append({
                    'arquivo': str(caminho_arquivo),
                    'tipo': 'PPTX',
                    'erro': error
                })
                self.stats['arquivos_corrompidos'] += 1

    def processar_diretorio(self):
        """Processa todos os arquivos do diretório"""
        self.logger.info(f"Iniciando análise do diretório: {self.diretorio_consolidado}")
        
        # Lista todos os arquivos (exceto subdiretórios que criamos)
        arquivos = []
        for arquivo in self.diretorio_consolidado.iterdir():
            if arquivo.is_file() and arquivo.name not in ['analisador_correcao_arquivos.py', 'consolidador_arquivos.py', 'requirements.txt']:
                if not arquivo.name.startswith('.'):  # Ignora arquivos ocultos
                    arquivos.append(arquivo)
        
        self.stats['total_arquivos'] = len(arquivos)
        self.logger.info(f"Total de arquivos para analisar: {len(arquivos)}")
        
        # Analisa cada arquivo
        for arquivo in arquivos:
            try:
                self.analisar_arquivo(arquivo)
            except Exception as e:
                self.logger.error(f"Erro ao analisar {arquivo.name}: {str(e)}")
                self.stats['erros'].append(f"Erro ao analisar {arquivo.name}: {str(e)}")

    def corrigir_arquivos_corrompidos(self):
        """Tenta corrigir arquivos corrompidos"""
        self.logger.info("Iniciando correção de arquivos corrompidos...")
        
        for arquivo_info in self.arquivos_corrompidos:
            caminho_arquivo = Path(arquivo_info['arquivo'])
            tipo = arquivo_info['tipo']
            
            try:
                if tipo == 'DOCX':
                    sucesso, resultado = self.tentar_reparar_docx(caminho_arquivo)
                    if sucesso:
                        self.stats['arquivos_corrigidos'] += 1
                        self.logger.info(f"Arquivo corrigido: {resultado}")
                    else:
                        self.logger.error(f"Não foi possível corrigir: {caminho_arquivo.name} - {resultado}")
                
                # Para outros tipos, implementar métodos de reparo específicos
                # elif tipo == 'XLSX':
                #     # Implementar reparo de XLSX
                # elif tipo == 'PDF':
                #     # Implementar reparo de PDF
                # elif tipo == 'PPTX':
                #     # Implementar reparo de PPTX
                
            except Exception as e:
                self.logger.error(f"Erro ao tentar corrigir {caminho_arquivo.name}: {str(e)}")
                self.stats['erros'].append(f"Erro ao corrigir {caminho_arquivo.name}: {str(e)}")

    def converter_xlsx_para_csv_todos(self):
        """Converte todos os arquivos XLSX para CSV"""
        self.logger.info("Iniciando conversão de arquivos XLSX para CSV...")
        
        for arquivo_xlsx in self.arquivos_xlsx:
            try:
                sucesso, resultado = self.converter_xlsx_para_csv(arquivo_xlsx)
                if sucesso:
                    self.stats['xlsx_convertidos'] += 1
                else:
                    self.stats['erros'].append(f"Erro ao converter {arquivo_xlsx.name}: {resultado}")
            except Exception as e:
                self.logger.error(f"Erro ao converter {arquivo_xlsx.name}: {str(e)}")
                self.stats['erros'].append(f"Erro ao converter {arquivo_xlsx.name}: {str(e)}")

    def gerar_relatorio(self):
        """Gera relatório final"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        relatorio_file = self.diretorio_logs / f"relatorio_analise_correcao_{timestamp}.json"
        
        relatorio = {
            'timestamp': timestamp,
            'estatisticas': self.stats,
            'arquivos_corrompidos': self.arquivos_corrompidos,
            'arquivos_xlsx_encontrados': [str(f) for f in self.arquivos_xlsx]
        }
        
        with open(relatorio_file, 'w', encoding='utf-8') as f:
            json.dump(relatorio, f, indent=2, ensure_ascii=False)
        
        self.logger.info(f"Relatório salvo em: {relatorio_file}")
        
        # Exibe resumo
        print("\n" + "="*60)
        print("RELATÓRIO FINAL - ANÁLISE E CORREÇÃO DE ARQUIVOS")
        print("="*60)
        print(f"Total de arquivos analisados: {self.stats['total_arquivos']}")
        print(f"Arquivos corrompidos encontrados: {self.stats['arquivos_corrompidos']}")
        print(f"Arquivos corrigidos: {self.stats['arquivos_corrigidos']}")
        print(f"Arquivos XLSX convertidos para CSV: {self.stats['xlsx_convertidos']}")
        print(f"Erros encontrados: {len(self.stats['erros'])}")
        
        if self.arquivos_corrompidos:
            print("\nARQUIVOS CORROMPIDOS:")
            for arquivo in self.arquivos_corrompidos:
                print(f"  - {Path(arquivo['arquivo']).name} ({arquivo['tipo']}): {arquivo['erro']}")
        
        if self.stats['erros']:
            print("\nERROS:")
            for erro in self.stats['erros']:
                print(f"  - {erro}")
        
        print("="*60)

    def executar(self):
        """Executa todo o processo"""
        try:
            self.logger.info("Iniciando análise e correção de arquivos...")
            
            # 1. Processar diretório
            self.processar_diretorio()
            
            # 2. Corrigir arquivos corrompidos
            if self.arquivos_corrompidos:
                self.corrigir_arquivos_corrompidos()
            
            # 3. Converter XLSX para CSV
            if self.arquivos_xlsx:
                self.converter_xlsx_para_csv_todos()
            
            # 4. Gerar relatório
            self.gerar_relatorio()
            
            self.logger.info("Processo concluído com sucesso!")
            
        except Exception as e:
            self.logger.error(f"Erro durante execução: {str(e)}")
            raise

def main():
    """Função principal"""
    diretorio_consolidado = "/Users/marcosdaniels/Downloads/5. FULL SALES 2/Consolidado"
    
    if not os.path.exists(diretorio_consolidado):
        print(f"ERRO: Diretório não encontrado: {diretorio_consolidado}")
        sys.exit(1)
    
    analisador = AnalisadorCorrecaoArquivos(diretorio_consolidado)
    analisador.executar()

if __name__ == "__main__":
    main()