#!/usr/bin/env python3
"""
Script para converter arquivos de diferentes formatos para TXT
Suporta: MD, CSV, PY, PDF, DOCX, VB, STY, NUMBERS e arquivos sem extensão
"""

import os
import shutil
import subprocess
import sys
from pathlib import Path

def instalar_dependencias():
    """Instala as dependências necessárias"""
    dependencias = [
        'PyPDF2',
        'python-docx',
        'pandas',
        'openpyxl'
    ]
    
    for dep in dependencias:
        try:
            if dep == 'PyPDF2':
                __import__('PyPDF2')
            elif dep == 'python-docx':
                __import__('docx')
            elif dep == 'pandas':
                __import__('pandas')
            elif dep == 'openpyxl':
                __import__('openpyxl')
        except ImportError:
            print(f"Instalando {dep}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", dep])

# Instalar dependências
instalar_dependencias()

# Variáveis de controle de disponibilidade
PDF_DISPONIVEL = False
DOCX_DISPONIVEL = False
PANDAS_DISPONIVEL = False

# Importações condicionais
try:
    import PyPDF2
    PDF_DISPONIVEL = True
except ImportError:
    PyPDF2 = None
    print("Aviso: PyPDF2 não disponível - conversão de PDF desabilitada")

try:
    from docx import Document
    DOCX_DISPONIVEL = True
except ImportError:
    Document = None
    print("Aviso: python-docx não disponível - conversão de DOCX desabilitada")

try:
    import pandas
    PANDAS_DISPONIVEL = True
except ImportError:
    pandas = None
    print("Aviso: pandas não disponível - algumas conversões podem falhar")

def converter_md_para_txt(arquivo_origem, arquivo_destino):
    """Converte arquivo Markdown para TXT"""
    try:
        with open(arquivo_origem, 'r', encoding='utf-8') as f:
            conteudo = f.read()
        
        with open(arquivo_destino, 'w', encoding='utf-8') as f:
            f.write(conteudo)
        return True
    except Exception as e:
        print(f"❌ Erro ao converter {arquivo_origem}: {e}")
        return False

def converter_csv_para_txt(arquivo_origem, arquivo_destino):
    """Converte arquivo CSV para TXT"""
    try:
        import pandas as pd
        df = pd.read_csv(arquivo_origem, encoding='utf-8')
        with open(arquivo_destino, 'w', encoding='utf-8') as f:
            f.write(df.to_string(index=False))
        return True
    except Exception as e:
        try:
            # Fallback: leitura simples
            with open(arquivo_origem, 'r', encoding='utf-8') as f:
                conteudo = f.read()
            with open(arquivo_destino, 'w', encoding='utf-8') as f:
                f.write(conteudo)
            return True
        except Exception as e2:
            print(f"❌ Erro ao converter {arquivo_origem}: {e2}")
            return False

def converter_py_para_txt(arquivo_origem, arquivo_destino):
    """Converte arquivo Python para TXT"""
    try:
        with open(arquivo_origem, 'r', encoding='utf-8') as f:
            conteudo = f.read()
        
        with open(arquivo_destino, 'w', encoding='utf-8') as f:
            f.write(f"# Código Python - {os.path.basename(arquivo_origem)}\n\n")
            f.write(conteudo)
        return True
    except Exception as e:
        print(f"❌ Erro ao converter {arquivo_origem}: {e}")
        return False

def converter_pdf_para_txt(arquivo_origem, arquivo_destino):
    """Converte arquivo PDF para TXT"""
    if not PDF_DISPONIVEL:
        raise ImportError("PyPDF2 não está disponível")
    
    try:
        with open(arquivo_origem, 'rb') as arquivo:
            leitor = PyPDF2.PdfReader(arquivo)
            texto = ""
            for pagina in leitor.pages:
                texto += pagina.extract_text() + "\n"
        
        with open(arquivo_destino, 'w', encoding='utf-8') as arquivo:
            arquivo.write(f"=== ARQUIVO PDF CONVERTIDO ===\n")
            arquivo.write(f"Origem: {arquivo_origem}\n")
            arquivo.write(f"Páginas: {len(leitor.pages)}\n")
            arquivo.write("=" * 50 + "\n\n")
            arquivo.write(texto)
        
        return True
    except Exception as e:
        print(f"Erro ao converter PDF {arquivo_origem}: {e}")
        return False

def converter_docx_para_txt(arquivo_origem, arquivo_destino):
    """Converte arquivo DOCX para TXT"""
    if not DOCX_DISPONIVEL:
        raise ImportError("python-docx não está disponível")
    
    try:
        doc = Document(arquivo_origem)
        texto = ""
        for paragrafo in doc.paragraphs:
            texto += paragrafo.text + "\n"
        
        with open(arquivo_destino, 'w', encoding='utf-8') as arquivo:
            arquivo.write(f"=== ARQUIVO DOCX CONVERTIDO ===\n")
            arquivo.write(f"Origem: {arquivo_origem}\n")
            arquivo.write(f"Parágrafos: {len(doc.paragraphs)}\n")
            arquivo.write("=" * 50 + "\n\n")
            arquivo.write(texto)
        
        return True
    except Exception as e:
        print(f"Erro ao converter DOCX {arquivo_origem}: {e}")
        return False

def converter_arquivo_generico(arquivo_origem, arquivo_destino):
    """Converte arquivos genéricos (VB, STY, sem extensão) para TXT"""
    try:
        with open(arquivo_origem, 'r', encoding='utf-8') as f:
            conteudo = f.read()
        
        with open(arquivo_destino, 'w', encoding='utf-8') as f:
            f.write(f"# Arquivo convertido - {os.path.basename(arquivo_origem)}\n\n")
            f.write(conteudo)
        return True
    except Exception as e:
        try:
            # Tentativa com encoding latin-1
            with open(arquivo_origem, 'r', encoding='latin-1') as f:
                conteudo = f.read()
            
            with open(arquivo_destino, 'w', encoding='utf-8') as f:
                f.write(f"# Arquivo convertido - {os.path.basename(arquivo_origem)}\n\n")
                f.write(conteudo)
            return True
        except Exception as e2:
            print(f"❌ Erro ao converter {arquivo_origem}: {e2}")
            return False

def converter_numbers_para_txt(arquivo_origem, arquivo_destino):
    """Converte arquivo Numbers para TXT (informativo)"""
    try:
        with open(arquivo_destino, 'w', encoding='utf-8') as f:
            f.write(f"# Arquivo Numbers - {os.path.basename(arquivo_origem)}\n\n")
            f.write("NOTA: Este é um arquivo Numbers (.numbers) que não pode ser convertido diretamente.\n")
            f.write("Para acessar o conteúdo, abra o arquivo no Numbers (macOS) ou exporte para CSV/Excel.\n")
            f.write(f"Localização original: {arquivo_origem}\n")
        return True
    except Exception as e:
        print(f"❌ Erro ao criar arquivo informativo para {arquivo_origem}: {e}")
        return False

def main():
    """Função principal"""
    print("🔄 Iniciando conversão de arquivos para TXT...")
    
    # Instalar dependências
    instalar_dependencias()
    
    pasta_origem = "Arquivos_Consolidados"
    pasta_destino = "Arquivos_TXT"
    
    # Criar pasta de destino
    if os.path.exists(pasta_destino):
        shutil.rmtree(pasta_destino)
    os.makedirs(pasta_destino)
    
    # Contadores
    total_arquivos = 0
    convertidos_sucesso = 0
    convertidos_erro = 0
    
    # Mapear conversões
    conversores = {
        '.md': converter_md_para_txt,
        '.csv': converter_csv_para_txt,
        '.py': converter_py_para_txt,
        '.pdf': converter_pdf_para_txt,
        '.docx': converter_docx_para_txt,
        '.vb': converter_arquivo_generico,
        '.sty': converter_arquivo_generico,
        '.txt': converter_arquivo_generico,
        '.numbers': converter_numbers_para_txt,
        '': converter_arquivo_generico  # Arquivos sem extensão
    }
    
    # Processar todos os arquivos
    for arquivo in os.listdir(pasta_origem):
        if arquivo.startswith('.'):
            continue
            
        caminho_origem = os.path.join(pasta_origem, arquivo)
        
        if not os.path.isfile(caminho_origem):
            continue
            
        total_arquivos += 1
        
        # Determinar extensão
        extensao = Path(arquivo).suffix.lower()
        if not extensao and not arquivo.endswith('.txt'):
            extensao = ''  # Arquivo sem extensão
        
        # Nome do arquivo de destino
        nome_base = Path(arquivo).stem if extensao else arquivo
        arquivo_destino = os.path.join(pasta_destino, f"{nome_base}.txt")
        
        print(f"📄 Convertendo: {arquivo} -> {nome_base}.txt")
        
        # Converter arquivo
        if extensao in conversores:
            if conversores[extensao](caminho_origem, arquivo_destino):
                convertidos_sucesso += 1
            else:
                convertidos_erro += 1
        else:
            # Tentar conversão genérica
            if converter_arquivo_generico(caminho_origem, arquivo_destino):
                convertidos_sucesso += 1
            else:
                convertidos_erro += 1
    
    # Criar arquivo de relatório
    relatorio_path = os.path.join(pasta_destino, "000_RELATORIO_CONVERSAO.txt")
    with open(relatorio_path, 'w', encoding='utf-8') as f:
        f.write("RELATÓRIO DE CONVERSÃO PARA TXT\n")
        f.write("=" * 50 + "\n\n")
        f.write(f"Total de arquivos processados: {total_arquivos}\n")
        f.write(f"Conversões bem-sucedidas: {convertidos_sucesso}\n")
        f.write(f"Conversões com erro: {convertidos_erro}\n")
        f.write(f"Taxa de sucesso: {(convertidos_sucesso/total_arquivos*100):.1f}%\n\n")
        f.write("TIPOS DE ARQUIVO CONVERTIDOS:\n")
        f.write("- Markdown (.md) -> TXT\n")
        f.write("- CSV (.csv) -> TXT (formatado como tabela)\n")
        f.write("- Python (.py) -> TXT (com cabeçalho)\n")
        f.write("- PDF (.pdf) -> TXT (texto extraído)\n")
        f.write("- Word (.docx) -> TXT (texto extraído)\n")
        f.write("- Visual Basic (.vb) -> TXT\n")
        f.write("- Style (.sty) -> TXT\n")
        f.write("- Numbers (.numbers) -> TXT (arquivo informativo)\n")
        f.write("- Arquivos sem extensão -> TXT\n")
        f.write("- Arquivos TXT -> TXT (cópia)\n")
    
    print(f"\n✅ Conversão concluída!")
    print(f"📊 Estatísticas:")
    print(f"   • Total de arquivos: {total_arquivos}")
    print(f"   • Convertidos com sucesso: {convertidos_sucesso}")
    print(f"   • Erros: {convertidos_erro}")
    print(f"   • Taxa de sucesso: {(convertidos_sucesso/total_arquivos*100):.1f}%")
    print(f"📁 Arquivos TXT salvos em: {pasta_destino}/")

if __name__ == "__main__":
    main()